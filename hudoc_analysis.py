import eel

import requests
import urllib.parse
from bs4 import BeautifulSoup as BS


from tqdm import tqdm
import os, json, shutil
import re
import time
from datetime import datetime
import sqlite3
from urllib.parse import urlencode

from docx import Document


eel.init("frontend")

for d in ["sets", "keywords", "docx"]:
    if not os.path.isdir(d):
        os.mkdir(d)

def get_html(url, headers):
    """get requests.get().text with retry on error
    :param url: url for requests.get
    :type url: str
    :param headers: headers for requests.get
    :type headers: dict
    :return: HTML response
    :rtype: str
    """

    try:
        return requests.get(url, headers=headers).text
    except:
        time.sleep(5)
        return get_html(url, headers)

echrcaselaw_headers = {
    "Host": "www.echrcaselaw.com",
    "User-Agent": "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:109.0) Gecko/20100101 Firefox/117.0",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
    "Accept-Language": "de,en-US;q=0.7,en;q=0.3",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "cross-site",
    "TE": "trailers"
}

def get_blog_urls(id, appno, accepted_blogs):
    """search for id on various blogs, gather all post links
    :param id: HUDOC itemid to search for
    :type id: str
    :param appno: First appno of document
    :type appno: str
    :param accepted_blogs: provide links to various blog posts mentioning each document, defaults to not including any blogs
    :type accepted_blogs: dict, optional
    :return: List of post links joined by ;;;
    :rtype: str
    """

    try:
        urls = []
        if accepted_blogs["strasbourg"]["include"]:
            bs = BS(requests.get(f"https://strasbourgobservers.com/?s={id}").text, "lxml")
            urls += list(set([u["href"] for u in bs.find_all("a", {"href": True}) if re.search("https\:\/\/strasbourgobservers.com\/\d+", u["href"])]))

        if accepted_blogs["verfassungsblog"]["include"]:
            bs = BS(requests.get(f"https://verfassungsblog.de/?s={id}").text, "lxml")
            urls += list(set([u["href"] for u in bs.find_all("a", {"href": True}, string="Continue reading >>")]))
        
        if accepted_blogs["voelkerrechtsblog"]["include"]:
            bs = BS(requests.get(f"https://voelkerrechtsblog.org/?s={id}").text, "lxml")
            urls += list(set([u["href"] for u in bs.find("div", {"class": "results-items"}).find_all("a", {"href": True})])) if bs.find("div", {"class": "results-items"}) else []
        
        if accepted_blogs["echrcaselaw"]["include"]:
            bs = BS(requests.get(f"https://www.echrcaselaw.com/en/?{urlencode({'s': appno})}", headers=echrcaselaw_headers).text, "lxml")
            urls += list(set([u["href"] for u in bs.find_all("a", {"class": "readmore", "href": True})]))

        return ";;;".join(urls) #sqlite does not support list fields
    except Exception as error:
        print(error)
        return ""

@eel.expose
def dl_hudoc(from_date, to_date, save_name, accepted_types = ["JUDGMENTS", "DECISIONS, ADVISORYOPINIONS"], accepted_langs = ["ENG"], respondents = [], accepted_blogs = {"strasbourg": False, "verfassungsblog": False, "voelkerrechtsblog": False, "echrcaselaw": False}, experimental_short = False, custom_query = "", base_set_name = ""):
    """main function to download metadata and HTML from HUDOC and parse text sections from HTML
    :param from_date: begin of download timespan in format yyyy-mm-dd
    :type from_date: str
    :param to_date: end of download timespan in format yyyy-mm-dd
    :type to_date: str
    :param save_name: filename for .db file, function cleans invalid characters, adds .db
    :type save_name: str
    :param accepted_types: doctypes to download from HUDOC as used in query param, defaults to JUDGEMENTS, DECISIONS, ADVISORYOPINIONS
    :type accepted_types: list, optional
    :param accepted_langs: only save docs in these languages, ISO codes, defaults to ENG
    :type accepted_langs: list, optional
    :param respondents: if not empty, only save docs with these respondents, Country ISO Code or english name, ideally provide both, defaults to all respondents
    :type respondents: list, optional
    :param accepted_blogs: provide links to various blog posts mentioning each document, defaults to not including any blogs
    :type accepted_blogs: dict, optional
    :param experimental_short: Filter out parties' submissions by formatting and keywords, may filter out court's assessment in some instances -> use at own risk, defaults to False
    :type experimental_short: bool, optional
    :param custom_query: if provided, overwrites other params and query formatting in favor of own query, recommended only for advanced users looking for additional filters, defaults to no custom query
    :type custom_query: str, optional
    :param base_set_name: if provided, copies base set and adds new downloads, useful when downloading with different filters for different timespans, defaults to no base set
    :type base_set_name: str, optional
    """
    
    try:
        if os.path.isfile(os.path.join("sets", f"{save_name}.db")):
            eel.setProgress(-1, "Filename already exists!")
            return
        if datetime.strptime(from_date, "%Y-%m-%d") > datetime.strptime(to_date, "%Y-%m-%d"):
            eel.setProgress(-1, "Start date > end date!")
            return
        
        save_name_match = re.search(r"[A-Za-z0-9\-\_]+", save_name)
        if not save_name_match:
            eel.setProgress(-1, "Invalid save name!")
            return
        save_name = save_name_match.group()

        HEADERS = {
            "Host": "hudoc.echr.coe.int",
            "User-Agent": "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:109.0) Gecko/20100101 Firefox/117.0",
            "Accept": "*/*",
            "Accept-Language": "de,en-US;q=0.7,en;q=0.3",
            "Accept-Encoding": "gzip, deflate, br",
            "Authorization": "Bearer",
            "Connection": "keep-alive",
            "Referer": "https://hudoc.echr.coe.int/Javascript/Shared/webWorker.js",
            "Sec-Fetch-Dest": "empty",
            "Sec-Fetch-Mode": "cors",
            "Sec-Fetch-Site": "same-origin",
            "TE": "trailers"
        }

        CHUNK_SIZE = 2000
        from_year = int(from_date[:from_date.find("-")])
        to_year = int(to_date[:to_date.find("-")])
        print(from_date, to_date, from_year, to_year)

        results = []
        for ind, year in enumerate(range(from_year, to_year + 1)): #HUDOC maxes at 10.000 results per query, independant of start and length -> chunking queries by year
            eel.setProgress(int(ind * 100.0 / max((to_year - from_year), 1)), f"Downloading metadata for {year}")
            params = {
                "query": custom_query if custom_query != "" else f'contentsitename:ECHR AND (NOT (doctype=PR OR doctype=HFCOMOLD OR doctype=HECOMOLD)){"AND (" + " OR ".join([f"""(documentcollectionid="{t}")""" for t in accepted_types]) + ")" if len(accepted_types) > 0 else ""} AND (kpdate>="{from_date if year == from_year else f"{year}-01-01"}T00:00:00.0Z" AND kpdate<="{to_date if year == to_year else f"{year}-12-31"}T00:00:00.0Z")',
                "select": "itemid,docname,doctype,appno,conclusion,importance,originatingbody,kpdate,kpdateastext,documentcollectionid,documentcollectionid2,languageisocode,advopidentifier,advopstatus,ecli,isplaceholder,separateopinion,article,respondent",
                "sort": "",
                "start": 0,
                "length": CHUNK_SIZE,
                "rankingModelId": "11111111-0000-0000-0000-000000000000"
            }

            query_url = f"https://hudoc.echr.coe.int/app/query/results?{urllib.parse.urlencode(params)}"

            print(f"Fetching {query_url}...")
            init_result = requests.get(query_url, headers=HEADERS).json()
            year_results = [r["columns"] for r in init_result["results"]]
            
            num_results = init_result["resultcount"]
            while len(year_results) < min(num_results,9999):
                add_url = query_url.replace("start=0", f"start={len(year_results)}")

                print(f"Got {len(year_results)}/{num_results}...")
                print(f"Fetching {add_url}...")
                new_results = requests.get(add_url, headers=HEADERS).json()
                year_results += [r["columns"] for r in new_results["results"]]
                if len(new_results["results"]) == 0:
                    break
            results += year_results
        
        respondent_match = f"v. (?:{'|'.join([f'(?:{r})' for r in respondents])})"
        results = [r for r in results if (not r["isplaceholder"] == "True") and r["languageisocode"] in accepted_langs and (len(respondents) == 0 or r["respondent"] in respondents or re.search(respondent_match, r["docname"]))]
        results = {r["itemid"]: r for r in results}
        fail_count = 0
        if base_set_name != "":
            shutil.copyfile(os.path.join("sets", base_set_name), os.path.join("sets",  f"{save_name}.db"))
        con = sqlite3.connect(os.path.join("sets", f"{save_name}.db"))
        cur = con.cursor()
        for ind, k in enumerate(tqdm(results.keys())):
            r = results[k]
            eel.setProgress(int(ind * 100.0 / len(results.keys())), f"Lade Dokument {ind + 1}/{len(results.keys())}")
            r["approve_status"] = 0
            result_url = f"https://hudoc.echr.coe.int/app/conversion/docx/html/body?library=ECHR&id={r['itemid']}"
            r["blog_urls"] = get_blog_urls(r["itemid"], r["appno"].split(";")[0], accepted_blogs)
            r["html"] = get_html(result_url, HEADERS)
            if r["html"] == "":
                fail_count += 1
                continue
            r.update(get_doc_info(r, experimental_short))

            r = dict(sorted(r.items())) #HUDOC sends metadata in varying key order

            if ind == 0 and base_set_name == "":
                cur.execute(f"CREATE TABLE docs({', '.join(r.keys())})")
            cur.execute(f"INSERT INTO docs VALUES ({', '.join([':' + k for k in r.keys()])})", r)
        
        con.commit()
        eel.setProgress(100, "Done")
        print(f"Fail count: {fail_count}")
    except Exception as error:
        print(error)
        eel.setProgress(-1, "Unhandled error during download, please contact developer!")

@eel.expose
def get_sets():
    """get list of .db files in sets dir and .json files in keywords dir
    :return: list containing two lists of filename strings [sets, keywords]
    :rtype: list
    """
    try:
        return [os.listdir("sets"), os.listdir("keywords")]
    except Exception as error:
        print(error)
        return []

@eel.expose
def read_kw_file(fn):
    """read .json file of given filename and return content as list of dicts
    :return: list containing keyword dicts, format [{"header": ["keyword1", "keyword2"], ...}, ...], keywords may start with regex:
    :rtype: list
    """
    try:
        with open(os.path.join("keywords", fn)) as f:
            return json.loads(f.read())
    except Exception as error:
        print(error)
        return []
    
@eel.expose
def set_doc_status(itemid, status):
    """update approval status of doc and save to .db
    :param itemid: HUDOC itemid of doc
    :type itemid: str
    :param status: approval status code with key 0 (undecided), 1 (approved), 2(disapproved)
    """
    try:
        cur.execute(f"UPDATE docs SET approve_status={status} WHERE itemid={itemid}")
        con.commit()
    except Exception as error:
        print(error)

@eel.expose
def save_keywords(keywords, fn):
    """save keyword file to .json
    :param keywords: list containing keyword dicts, format [{"header": ["keyword1", "keyword2"], ...}, ...], keywords may start with regex:
    :type keywords: list
    :param fn: filename within keywords dir
    :type fn: str
    """
    try:
        with open(os.path.join("keywords", fn), "w", encoding="utf-8") as f:
            json.dump(keywords, f, ensure_ascii=False, indent=4)
    except Exception as error:
        print(error)

def spans_to_text(spans):
    """parse HUDOC html so that spans of same p are grouped together
    :param spans: list of bs4 span objects
    :type spans: list
    :return: merged text, paragraphs separated by \\n
    :rtype: str
    """
    try:
        text = ""
        par = ""
        parent_tag = BS().new_tag("b")
        for span in spans:
            if (span.parent if span.parent.name == "p" else span.parent.parent) == parent_tag:
                par += span.text
            else:
                text += par.replace("\n", "") + "\n\n"
                parent_tag = span.parent if span.parent.name == "p" else span.parent.parent
                par = span.text
        while re.search("\s*\n\s*\n\s*\n\s*", text):
            text = re.sub("\s*\n\s*\n\s*\n\s*", "\n\n", text)
        return text.strip()
    except Exception as error:
        print(error)
        return ""

def get_texts(bs, doctype, experimental_short = True, lang = "ENG"):
    """extract violation, fact and law sections from BS object of HUDOC HTML
    :param bs: BeautifulSoup representation of HUDOC HTML
    :type bs: :class: 'bs4.BeautifulSoup'
    :param experimental_short: Filter out parties' submissions by formatting and keywords, may filter out court's assessment in some instances -> use at own risk, defaults to False
    :type experimental_short: bool, optional
    :param lang: document language ISO code, defaults to english
    :type lang: str
    :return: violation text if found else empty string, fact text, law text
    :rtype: str, str, str
    """
    try:
        if "CLIN" in doctype:
            violation_text = ""
            fact_text = ""
            law_text = ""
            for p in bs.find_all("p", {"class": True}):
                for span in p.find_all("span"):
                    t = span.text.strip()
                    if t in ["violation", "no violation"]:
                        if violation_text != "":
                            violation_text += "\n"
                        violation_text += p.text
                    elif t == "Facts":
                        content_ps = [tag.text.strip() for tag in bs.find_all("p", {"class": p["class"]})]
                        cutoff = 1
                        for ind, cp in enumerate(content_ps):
                            if cp.startswith("Facts -"):
                                cutoff = ind
                        fact_text = "\n\n".join(content_ps[:cutoff + 1])[len("Facts - "):]
                        law_text = "\n\n".join(content_ps[cutoff + 1:])[len("Law - "):]
            return violation_text, fact_text, law_text
        
        if "JUDGMENTS" in doctype or "DECISION" in doctype:
            law = False
            keep = True
            fact_index = 0
            law_index = 0
            spans = bs.find_all("span")
            for ind, span in enumerate(spans):
                if span.text.strip().lower() in ["the facts", "subject matter of the case"]:
                    fact_index = ind
                if span.text.strip().lower() in ["the law", "the court’s assessment"]:
                    law = True
                    law_index = ind
                if not law:
                    continue
                if span.text.strip().lower() == "the parties’ submissions":
                    keep = False
                elif span.text.strip().lower().startswith("the government"):
                    keep = False
                elif span.text.strip().lower().startswith("the applicant"):
                    keep = False
                elif span.text.strip().lower().startswith("the court"):
                    keep = True
                if experimental_short and lang == "ENG" and not keep:
                    span.string = ""
            
            fact_text = spans_to_text(spans[fact_index:law_index])        
            law_text = spans_to_text(spans[law_index:])

            return "", fact_text[:2000], law_text

        return "", "", spans_to_text(bs.find_all("span"))
    except Exception as error:
        print(error)
        return "", "", ""

def get_opinion_header(html):
    """extract header for separate opinion from HUDOC HTML, assumes that separate opinion exists in doc (check with metadata separateopinion)
    :param html: HUDOC HTML string
    :type html: str
    :return: separate opinion header, falls back to Yes
    :rtype: str
    """
    try:
        for span in BS(html, "lxml").find_all("span"):
            if "OPINION" in span.text:
                return span.text
        return "Yes"
    except Exception as error:
        print(error)
        return "Yes"

def get_doc_info(h, experimental_short):
    """gather output formatted metadata and text sections from metadata and HTML returned by HUDOC
    :param h: dict with metadata from HUDOC results endpoint and additional "html" field with HUDOC HTML string value as constructed by dl_hudoc
    :type h: dict
    :param experimental_short: Filter out parties' submissions by formatting and keywords, may filter out court's assessment in some instances -> use at own risk, defaults to False
    :type experimental_short: bool, optional
    :return: dict with new string fields header, url, conclusion_text, appnos (separated by \\n), date, separate (with header), articles (separated by \\n), violation_text
    :rtype: dict
    """
    try:
        bs = BS(h["html"], "lxml")
        info = {}
        info["header"] = f"G{h['importance']}: {h['docname']}"
        info["url"] = f"https://hudoc.echr.coe.int/eng?i={h['itemid']}"
        
        info["conclusion_text"] = h["conclusion"]
        replace_ind = [m.start() for m in re.finditer(";\S", info["conclusion_text"])]
        info["conclusion_text"] = list(info["conclusion_text"])
        for i in replace_ind:
            info["conclusion_text"][i] = "\n"
        info["conclusion_text"] = "".join(info["conclusion_text"])

        info["appnos"] = "\n".join(h["appno"].split(";")[:3]) if "appno" in h.keys() else ""
        info["date"] = h["kpdateastext"][:h["kpdateastext"].find(" ")]
        info["separate"] = f"Separate Opinion(s): {'No' if not h['separateopinion'] or h['separateopinion'] in ['', 'FALSE'] else get_opinion_header(h['html'])}"
        info["articles"] = "Article(s)\n" + "\n".join(h["article"].split(";"))

        info["violation_text"], info["fact_text"], info["law_text"] = get_texts(bs, h["documentcollectionid"], experimental_short, h["languageisocode"])
        return info
    except Exception as error:
        print(error)
        return {
            "header": "",
            "url": "",
            "conclusion_text": "",
            "appnos": "",
            "date": "",
            "separate": "",
            "articles": "",
            "violation_text": ""
        }

@eel.expose
def get_docs(set_name, hide_seen, kw_dict):
    """get keyword and regex matches structured by header tree from document .db
    :param set_name: .db filename within set directory
    :type set_name: str
    :param hide_seen: if true, excludes docs from results where approval status is not 0 (undecided), to better structure work in longer projects
    :type hide_seen: bool
    :param kw_dict: list containing keyword dicts, format [{"header": ["keyword1", "keyword2"], ...}, ...], keywords may start with regex:
    :type kw_dict: list
    :return: dict containing all matching docs with HUDOC itemid as key, for fields see dl_hudoc, also returns header tree in format {article_num(str for Protocol articles): {"header": ["itemid1", "itemid2", ...], ...}, ...}
    :rtype: dict, dict
    """
    global con, cur
    try:
        keywords = []
        for h in kw_dict:
            keywords += kw_dict[h]
        keywords = list(set(keywords))
        re_keywords = [k for k in keywords if k.startswith("regex:")]
        keywords = [k for k in keywords if not k.startswith("regex:")]
        
        con = sqlite3.connect(os.path.join("sets", set_name))
        con.row_factory = sqlite3.Row
        cur = con.cursor()

        a_docs = {}
        for row in cur.execute("SELECT * FROM docs"):
            d = dict(row)
            if hide_seen and d["approve_status"] != 0:
                continue
            d["kw_count"] = {k: d["fact_text"].count(k) + d["law_text"].count(k) for k in keywords if k in d["fact_text"] or k in d["law_text"]}
            d["kw_count"].update({k: len(re.findall(k[len("regex:"):], d["fact_text"] + d["law_text"])) for k in re_keywords if re.search(k[len("regex:"):], d["fact_text"] + d["law_text"])})
            
            if len(d["kw_count"].keys()) > 0:
                for k in d["kw_count"].keys(): #highlight keywords
                    if k.startswith("regex"):
                        for res in re.findall(k[len("regex:"):], d["fact_text"]):
                            d["fact_text"] = d["fact_text"].replace(res, f'<a class="bg-yellow-200">{res}</a>')
                        for res in re.findall(k[len("regex:"):], d["law_text"]):
                            d["law_text"] = d["law_text"].replace(res, f'<a class="bg-yellow-200">{res}</a>')
                    else:
                        d["fact_text"] = d["fact_text"].replace(k, f'<a class="bg-yellow-200">{k}</a>')
                        d["law_text"] = d["law_text"].replace(k, f'<a class="bg-yellow-200">{k}</a>')
                
                a_docs[d["itemid"]] = d
        
        tree = {}
        
        articles = []
        for id in a_docs.keys():
            articles += a_docs[id]["article"].split(";")
        raw_articles = list(set(articles))
        articles = sorted([a for a in raw_articles if not "-" in a], key=lambda x: int(max(re.findall("\d+", x))) if re.search("\d+", x) else 0)
        articles += sorted([a for a in raw_articles if re.search("^P\d+\-\d+$", a)], key=lambda x: int(x[1:x.find("-")]) * 500 + int(x[x.find("-") + 1:]))

        for a in articles:
            hits = {id: a_docs[id] for id in a_docs.keys() if a in a_docs[id]["article"].split(";")}
            header_map = {}
            for h in kw_dict.keys():
                subhits = []
                for k in kw_dict[h]:
                    subhits += [(id, hits[id]["importance"]) for id in hits.keys() if k in hits[id]["kw_count"].keys()]
                if len(subhits) > 0:
                    header_map[h] = [hit[0] for hit in sorted(list(set(subhits)), key=lambda x: x[1])]
            if len(header_map.keys()) > 0:
                tree[a] = header_map

        return a_docs, tree
    except Exception as error:
        print(error)
        return {}, {}

@eel.expose
def export_docx(set_name, kw_dict, include_undecided = False):
    """write keyword matches from .db to docx-file
    :param set_name: .db filename within set directory, file generated by dl_hudoc()
    :type set_name: str
    :param kw_dict: list containing keyword dicts, format [{"header": ["keyword1", "keyword2"], ...}, ...], keywords may start with regex:
    :type kw_dict: list
    :param include_undecided: include docs with approval status 0 (undecided), defaults to not include undecided
    :type include_undecided: bool, optional
    """
    try:
        eel.setProgress(0, "Building header tree")
        docs, tree = get_docs(set_name, False, kw_dict)
        docs = {id: docs[id] for id in docs.keys() if docs[id]["approve_status"] == 1 or (include_undecided and docs[id]["approve_status"] == 0)}

        docx = Document()
        docx.add_heading("ECHR case law overview (all documents © ECHR-CEDH)", 0)
        for ind, a in enumerate(tqdm(tree.keys())):
            eel.setProgress(int(ind * 95.0 / len(tree.keys())), f"Exporting article {a}")
            docx.add_heading(f"Article {a[a.find('-') + 1:]} Protocol {a[1:a.find('-')]}" if a.startswith("P") else f"Article {a} ECHR", 1)
            
            for header in tree[a]:
                docx.add_heading(header, 2)
                for id in tree[a][header]:
                    if not id in docs.keys():
                        continue
                    h = docs[id]
                    docx.add_heading(h["header"], 3)
                    docx.add_paragraph(h["violation_text"] if h["violation_text"] != "" else h["conclusion_text"])
                    table = docx.add_table(rows=1, cols=3)

                    title_row = table.rows[0].cells
                    title_row[0].merge(title_row[1]).merge(title_row[2]).text = h["docname"]

                    url_row = table.add_row().cells
                    url_row[0].merge(url_row[1]).merge(url_row[2]).text = h["url"]

                    info_row = table.add_row().cells
                    info_row[0].text = h["appnos"]
                    info_row[1].text = h["date"]
                    info_row[2].text = h["separate"]

                    art_row = table.add_row().cells
                    art_row[0].text = h["articles"]
                    art_row[1].merge(art_row[2]).text = "Conclusion(s)\n" + "\n".join(re.split("\S;\S", h["conclusion_text"]))

                    fact_row = table.add_row().cells
                    fact_row[0].text = "Facts"
                    fact_row[1].merge(fact_row[2]).text = h["fact_text"]

                    law_row = table.add_row().cells
                    law_row[0].text = "Law"
                    law_row[1].merge(law_row[2]).text = h["law_text"]

        eel.setProgress(95, "Saving...")
        docx.save(os.path.join("docx", set_name.replace(".db", ".docx")))
        eel.setProgress(100, "Done")
    
    except Exception as error:
        print(error)
        eel.setProgress(-1, "Unhandled error during export, please contact developer!")

eel.start("frontend.html")